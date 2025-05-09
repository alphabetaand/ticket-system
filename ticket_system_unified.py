import os
#import tkinter as tk
#import customtkinter as ctk
from urllib.parse import urlparse
import cv2
import qrcode
import requests
import time
import json
from threading import Thread
from flask import Flask, request, jsonify, render_template_string, send_file
from passlib.context import CryptContext
from PIL import Image, ImageTk
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
from docx.shared import Pt
from datetime import datetime
from io import BytesIO
import webbrowser
import logging
from flask_cors import CORS
import numpy as np

# Configuration
QR_FOLDER = "qrcodes/"
CAMERA_INDEX = 0
ADMIN_PASSWORD = os.getenv('ADMIN_PASSWORD', 'admin123')  # À changer en production
FLASK_PORT = int(os.getenv('PORT', 5000))
MAX_HISTORY_ENTRIES = 50

# Configuration du logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialisation sécurité
pwd_context = CryptContext(
    schemes=["pbkdf2_sha256"],
    default="pbkdf2_sha256",
    pbkdf2_sha256__default_rounds=30000
)
ADMIN_PASSWORD_HASH = pwd_context.hash(ADMIN_PASSWORD)

# Template HTML pour l'interface mobile
MOBILE_TEMPLATE = """
<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Système de Tickets</title>
    <style>
        :root {
            --primary: #2b6cb0;
            --secondary: #4299e1;
            --success: #38a169;
            --danger: #e53e3e;
            --dark: #2d3748;
            --light: #f7fafc;
        }
        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        }
        body {
            background-color: #edf2f7;
            padding: 20px;
        }
        .container {
            max-width: 100%;
            margin: 0 auto;
            background: white;
            border-radius: 10px;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            overflow: hidden;
        }
        .header {
            background-color: var(--primary);
            color: white;
            padding: 15px;
            text-align: center;
        }
        .input-group {
            padding: 15px;
            border-bottom: 1px solid #e2e8f0;
        }
        input {
            width: 100%;
            padding: 12px;
            margin-bottom: 10px;
            border: 1px solid #cbd5e0;
            border-radius: 5px;
            font-size: 16px;
        }
        .button-group {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 10px;
            padding: 15px;
        }
        button {
            padding: 12px;
            border: none;
            border-radius: 5px;
            font-size: 16px;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.2s;
        }
        .btn-primary {
            background-color: var(--primary);
            color: white;
        }
        .btn-secondary {
            background-color: var(--secondary);
            color: white;
        }
        .btn-success {
            background-color: var(--success);
            color: white;
        }
        .btn-danger {
            background-color: var(--danger);
            color: white;
        }
        .status-container {
            padding: 15px;
        }
        .status {
            padding: 12px;
            border-radius: 5px;
            margin-bottom: 10px;
            display: none;
        }
        .status-valid {
            background-color: #c6f6d5;
            color: #22543d;
        }
        .status-invalid {
            background-color: #fed7d7;
            color: #742a2a;
        }
        .history {
            padding: 15px;
            max-height: 300px;
            overflow-y: auto;
        }
        .history-item {
            padding: 8px 0;
            border-bottom: 1px solid #e2e8f0;
            font-size: 14px;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>AGUIART & CO</h1>
        </div>
        
        <div class="input-group">
            <input type="number" id="ticketInput" placeholder="Numéro de ticket">
        </div>
        
        <div class="button-group">
            <button class="btn-success" id="validateBtn">Valider</button>
            <button class="btn-primary" id="verifyBtn">Vérifier</button>
        </div>
        
        <div class="button-group">
            <button class="btn-secondary" id="exportBtn">Exporter</button>
            <button class="btn-danger" id="adminBtn">Admin</button>
        </div>
        
        <div class="status-container">
            <div class="status" id="statusMessage"></div>
        </div>
        
        <div class="history" id="history"></div>
    </div>

    <script>
        const apiBaseUrl = window.location.origin;
        
        const ticketInput = document.getElementById('ticketInput');
        const validateBtn = document.getElementById('validateBtn');
        const verifyBtn = document.getElementById('verifyBtn');
        const exportBtn = document.getElementById('exportBtn');
        const adminBtn = document.getElementById('adminBtn');
        const statusMessage = document.getElementById('statusMessage');
        const historyDiv = document.getElementById('history');
        
        validateBtn.addEventListener('click', validateTicket);
        verifyBtn.addEventListener('click', verifyTicket);
        exportBtn.addEventListener('click', exportData);
        adminBtn.addEventListener('click', showAdminPanel);
        
        async function validateTicket() {
            const ticketNumber = ticketInput.value;
            if (!ticketNumber) return;
            
            try {
                const response = await fetch(`${apiBaseUrl}/validate`, {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({ ticket: ticketNumber })
                });
                
                const data = await response.json();
                showStatus(data.message || "Ticket validé avec succès", 'valid');
                addHistoryEntry(`Validé: Ticket ${ticketNumber}`);
            } catch (error) {
                showStatus("Erreur lors de la validation", 'invalid');
            }
        }
        
        async function verifyTicket() {
            const ticketNumber = ticketInput.value;
            if (!ticketNumber) return;
            
            try {
                const response = await fetch(`${apiBaseUrl}/verify?ticket=${ticketNumber}`);
                const data = await response.json();
                
                if (data.status === 'validé') {
                    showStatus(`Ticket ${ticketNumber} est valide`, 'valid');
                } else {
                    showStatus(`Ticket ${ticketNumber} est invalide`, 'invalid');
                }
                
                addHistoryEntry(`Vérifié: Ticket ${ticketNumber} (${data.status})`);
            } catch (error) {
                showStatus("Erreur lors de la vérification", 'invalid');
            }
        }
        
        async function exportData() {
            try {
                const response = await fetch(`${apiBaseUrl}/export_word`);
                if (response.ok) {
                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = 'tickets_valides.docx';
                    document.body.appendChild(a);
                    a.click();
                    document.body.removeChild(a);
                    showStatus("Export réussi", 'valid');
                } else {
                    showStatus("Erreur lors de l'export", 'invalid');
                }
            } catch (error) {
                showStatus("Erreur réseau", 'invalid');
            }
        }
        
        function showAdminPanel() {
            const password = prompt("Mot de passe admin:");
            if (password) {
                fetch(`${apiBaseUrl}/admin`, {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({ password: password })
                })
                .then(response => response.json())
                .then(data => {
                    if (data.success) {
                        showStatus("Accès admin autorisé", 'valid');
                        alert("Fonctionnalité admin - À implémenter");
                    } else {
                        showStatus("Accès refusé", 'invalid');
                    }
                })
                .catch(error => {
                    showStatus("Erreur de connexion", 'invalid');
                });
            }
        }
        
        function showStatus(message, type) {
            statusMessage.textContent = message;
            statusMessage.className = `status status-${type}`;
            statusMessage.style.display = 'block';
        }
        
        function addHistoryEntry(entry) {
            const entryElement = document.createElement('div');
            entryElement.className = 'history-item';
            entryElement.textContent = `${new Date().toLocaleTimeString()} - ${entry}`;
            historyDiv.prepend(entryElement);
        }
        
        function loadHistory() {
            fetch(`${apiBaseUrl}/history`)
                .then(response => response.json())
                .then(entries => {
                    entries.forEach(entry => {
                        const entryElement = document.createElement('div');
                        entryElement.className = 'history-item';
                        entryElement.textContent = entry;
                        historyDiv.appendChild(entryElement);
                    });
                })
                .catch(error => {
                    console.error("Erreur chargement historique:", error);
                });
        }
        
        loadHistory();
    </script>
</body>
</html>
"""

def get_db_connection():
    """Adapte automatiquement le connecteur à l'environnement"""
    if 'DATABASE_URL' in os.environ:
        # Mode production - PostgreSQL
        import psycopg2
        try:
            result = urlparse(os.getenv('DATABASE_URL'))
            conn = psycopg2.connect(
                dbname=result.path[1:],
                user=result.username,
                password=result.password,
                host=result.hostname,
                port=result.port
            )
            logger.info("Connecté à PostgreSQL")
            return conn
        except Exception as e:
            logger.error(f"Erreur PostgreSQL: {e}")
            raise
    else:
        # Mode développement - SQLite
        import sqlite3
        conn = sqlite3.connect('local_tickets.db')
        logger.info("Connecté à SQLite (local_tickets.db)")
        return conn

def init_db():
    """Initialise la structure de la base de données"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Script compatible SQLite et PostgreSQL
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS tickets (
                ticket_number INTEGER PRIMARY KEY,
                status TEXT DEFAULT 'invalide',
                qr_hash TEXT UNIQUE,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Adaptation pour SQLite
        if 'sqlite' in str(conn):
            cursor.execute("PRAGMA journal_mode=WAL")
        
        conn.commit()
        logger.info("Base de données initialisée")
    except Exception as e:
        logger.error(f"Erreur d'initialisation DB: {e}")
        raise
    finally:
        if 'conn' in locals():
            conn.close()

class QRScanner:
    def __init__(self, callback):
        self.cam = cv2.VideoCapture(CAMERA_INDEX)
        self.callback = callback
        self.running = False
        self.last_scanned = None
        self.last_scan_time = 0

    def start(self):
        if not self.cam.isOpened():
            logger.error("Impossible d'ouvrir la caméra")
            return False
        
        self.running = True
        Thread(target=self._scan_loop, daemon=True).start()
        return True

    def _scan_loop(self):
        while self.running:
            ret, frame = self.cam.read()
            if not ret:
                logger.warning("Erreur de lecture de la caméra")
                time.sleep(1)
                continue
                
            detector = cv2.QRCodeDetector()
            data, _, _ = detector.detectAndDecode(frame)
            
            if data and data != self.last_scanned or time.time() - self.last_scan_time > 5:
                self.last_scanned = data
                self.last_scan_time = time.time()
                self.callback(data)
                
            time.sleep(0.1)

    def stop(self):
        self.running = False
        if self.cam.isOpened():
            self.cam.release()
        logger.info("Scanner QR arrêté")

class TicketApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Ticket System")
        self.geometry("1000x800")
        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")
        
        self.scanner = None
        self._setup_ui()
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        
    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(3, weight=1)
        
        self.entry = ctk.CTkEntry(self, placeholder_text="Numéro de ticket")
        self.entry.grid(row=0, column=0, padx=20, pady=10, sticky="ew", columnspan=2)
        
        self.btn_validate = ctk.CTkButton(self, text="Valider", command=self._validate_ticket)
        self.btn_validate.grid(row=1, column=0, padx=10, pady=5, sticky="ew")
        
        self.btn_verify = ctk.CTkButton(self, text="Vérifier", command=self._verify_ticket)
        self.btn_verify.grid(row=1, column=1, padx=10, pady=5, sticky="ew")
        
        self.btn_scan = ctk.CTkButton(self, text="Scanner QR Code", command=self._toggle_scan)
        self.btn_scan.grid(row=2, column=0, padx=10, pady=5, columnspan=2, sticky="ew")
        
        self.cam_label = ctk.CTkLabel(self, text="Aucun flux vidéo")
        self.cam_label.grid(row=3, column=0, padx=10, pady=10, columnspan=2, sticky="nsew")
        
        self.history = ctk.CTkTextbox(self, wrap="word")
        self.history.grid(row=4, column=0, padx=10, pady=10, columnspan=2, sticky="nsew")
        
        self.btn_export = ctk.CTkButton(self, text="Exporter Word", command=self._export_word)
        self.btn_export.grid(row=5, column=0, padx=10, pady=5, sticky="ew")
        
        self.btn_admin = ctk.CTkButton(self, text="Admin", fg_color="#BF3A3A", command=self._show_admin)
        self.btn_admin.grid(row=5, column=1, padx=10, pady=5, sticky="ew")
        
        self.btn_mobile = ctk.CTkButton(self, text="Ouvrir Mobile", fg_color="#3A7EBF", command=self._open_mobile)
        self.btn_mobile.grid(row=6, column=0, padx=10, pady=5, columnspan=2, sticky="ew")
        
        self.status_label = ctk.CTkLabel(self, text="Prêt")
        self.status_label.grid(row=7, column=0, columnspan=2, sticky="ew")
    
    def _validate_ticket(self):
        ticket_num = self.entry.get()
        if not ticket_num.isdigit():
            self._update_history("Erreur: Numéro invalide")
            return
        
        try:
            response = requests.post(
                f'http://localhost:{FLASK_PORT}/validate',
                json={'ticket': ticket_num},
                timeout=5
            )
            response.raise_for_status()
            data = response.json()
            self._update_history(data.get('message', 'Ticket validé'))
            self._update_status("Validation réussie")
        except requests.RequestException as e:
            self._update_history(f"Erreur réseau: {str(e)}")
            self._update_status("Erreur de connexion")
        except Exception as e:
            self._update_history(f"Erreur: {str(e)}")
            self._update_status("Erreur inattendue")
    
    def _verify_ticket(self):
        ticket_num = self.entry.get()
        if not ticket_num:
            return
        
        try:
            response = requests.get(
                f'http://localhost:{FLASK_PORT}/verify?ticket={ticket_num}',
                timeout=5
            )
            response.raise_for_status()
            data = response.json()
            self._update_history(f"Ticket {data['ticket']}: {data['status']}")
            self._update_status("Vérification terminée")
        except requests.RequestException as e:
            self._update_history(f"Erreur réseau: {str(e)}")
            self._update_status("Erreur de connexion")
        except Exception as e:
            self._update_history(f"Erreur: {str(e)}")
            self._update_status("Erreur inattendue")
    
    def _toggle_scan(self):
        if self.scanner and self.scanner.running:
            self.scanner.stop()
            self.btn_scan.configure(text="Scanner QR Code")
            self.cam_label.configure(image=None, text="Scanner désactivé")
            self._update_status("Scanner arrêté")
        else:
            self.scanner = QRScanner(self._handle_qr_scan)
            if self.scanner.start():
                self.btn_scan.configure(text="Arrêter le scan")
                self._update_status("Scanner activé - Cherche QR codes...")
                self._update_camera_preview()
            else:
                self._update_history("Erreur: Impossible d'accéder à la caméra")
                self._update_status("Erreur caméra")
    
    def _update_camera_preview(self):
        if not self.scanner or not self.scanner.running:
            return
        
        try:
            ret, frame = self.scanner.cam.read()
            if ret:
                frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                img = Image.fromarray(frame)
                imgtk = ImageTk.PhotoImage(image=img)
                self.cam_label.configure(image=imgtk)
                self.cam_label.image = imgtk
                self.after(50, self._update_camera_preview)
        except Exception as e:
            logger.error(f"Erreur aperçu caméra: {e}")
    
    def _handle_qr_scan(self, data):
        try:
            ticket_num = data.split(':')[1]
            if not ticket_num.isdigit():
                raise ValueError("Format QR code invalide")
                
            self.entry.delete(0, tk.END)
            self.entry.insert(0, ticket_num)
            self._verify_ticket()
            self._update_status("QR code scanné")
        except Exception as e:
            self._update_history("QR Code invalide")
            self._update_status("Erreur scan QR")
    
    def _export_word(self):
        try:
            response = requests.get(
                f'http://localhost:{FLASK_PORT}/export_word',
                timeout=10
            )
            response.raise_for_status()
            
            file_path = os.path.join(os.getcwd(), 'tickets_valides.docx')
            with open(file_path, 'wb') as f:
                f.write(response.content)
            
            self._update_history("Export Word généré")
            self._update_status("Export réussi")
            
            if os.name == 'nt':
                os.startfile(file_path)
            else:
                webbrowser.open(file_path)
        except requests.RequestException as e:
            self._update_history(f"Erreur réseau: {str(e)}")
            self._update_status("Erreur export")
        except Exception as e:
            self._update_history(f"Erreur export: {str(e)}")
            self._update_status("Erreur export")
    
    def _show_admin(self):
        password = ctk.CTkInputDialog(
            text="Mot de passe admin:",
            title="Authentification"
        ).get_input()
        
        if password:
            try:
                response = requests.post(
                    f'http://localhost:{FLASK_PORT}/admin',
                    json={'password': password},
                    timeout=5
                )
                response.raise_for_status()
                
                if response.json().get('success'):
                    self._show_admin_dashboard()
                    self._update_status("Accès admin autorisé")
                else:
                    self._update_history("Accès refusé")
                    self._update_status("Mot de passe incorrect")
            except requests.RequestException as e:
                self._update_history(f"Erreur réseau: {str(e)}")
                self._update_status("Erreur connexion")
            except Exception as e:
                self._update_history(f"Erreur: {str(e)}")
                self._update_status("Erreur inattendue")
    
    def _show_admin_dashboard(self):
        admin_window = ctk.CTkToplevel(self)
        admin_window.title("Admin Dashboard")
        admin_window.geometry("800x600")
        
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            
            cursor.execute("SELECT COUNT(*) FROM tickets WHERE status='validé'")
            validated = cursor.fetchone()[0]
            cursor.execute("SELECT COUNT(*) FROM tickets")
            total = cursor.fetchone()[0]
            
            if total == 0:
                validated, total = 1, 1
                show_warning = True
            else:
                show_warning = False
            
            cursor.execute("""
                SELECT ticket_number, status, timestamp 
                FROM tickets 
                ORDER BY timestamp DESC 
                LIMIT 10
            """)
            recent_tickets = cursor.fetchall()
        
            admin_window.grid_columnconfigure(0, weight=1)
            admin_window.grid_rowconfigure(1, weight=1)
            
            stats_frame = ctk.CTkFrame(admin_window)
            stats_frame.grid(row=0, column=0, padx=20, pady=10, sticky="ew")
            
            ctk.CTkLabel(stats_frame, text=f"Tickets Validés: {validated}/{total}").pack(pady=5)
            
            fig, ax = plt.subplots(figsize=(6, 4))
            
            if show_warning:
                ax.text(0.5, 0.5, 'Aucune donnée disponible', 
                       ha='center', va='center')
                ax.set_title('Aucun ticket enregistré')
            else:
                ax.pie([validated, total-validated], 
                      labels=['Validés', 'Non validés'], 
                      autopct='%1.1f%%',
                      colors=['#2A8C36', '#BF3A3A'])
                ax.set_title('Statut des Tickets')
            
            chart_frame = ctk.CTkFrame(admin_window)
            chart_frame.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
            
            canvas = FigureCanvasTkAgg(fig, master=chart_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill='both', expand=True)
            
            tickets_frame = ctk.CTkFrame(admin_window)
            tickets_frame.grid(row=2, column=0, padx=10, pady=10, sticky="ew")
            
            ctk.CTkLabel(tickets_frame, text="Derniers tickets:").pack()
            
            for ticket in recent_tickets:
                status_color = "#2A8C36" if ticket[1] == 'validé' else "#BF3A3A"
                ticket_text = f"Ticket {ticket[0]} - {ticket[1]} - {ticket[2]}"
                label = ctk.CTkLabel(
                    tickets_frame, 
                    text=ticket_text,
                    text_color=status_color
                )
                label.pack(anchor='w')
            
            btn_frame = ctk.CTkFrame(admin_window)
            btn_frame.grid(row=3, column=0, padx=10, pady=10, sticky="ew")
            
            def on_delete_validated():
                try:
                    response = requests.get(f'http://localhost:{FLASK_PORT}/count_validated')
                    count = response.json().get("count", 0)
                    
                    if count == 0:
                        self._update_history("Aucun ticket validé à supprimer")
                        return
                    
                    confirm = ctk.CTkInputDialog(
                        text=f"⚠️ Supprimer {count} tickets validés ?\nMot de passe admin:",
                        title=f"Confirmation suppression"
                    ).get_input()
                    
                    if confirm:
                        response = requests.post(
                            f'http://localhost:{FLASK_PORT}/delete_validated',
                            json={'password': confirm},
                            timeout=5
                        )
                        data = response.json()
                        
                        if response.status_code == 200:
                            self._update_history(data["message"])
                            admin_window.destroy()
                            self._show_admin_dashboard()
                        else:
                            self._update_history(f"Erreur: {data.get('error', 'Inconnue')}")
                
                except Exception as e:
                    self._update_history(f"Échec suppression: {str(e)}")
            
            delete_btn = ctk.CTkButton(
                btn_frame,
                text="🗑️ Supprimer TOUS les tickets validés",
                fg_color="#ff5555",
                hover_color="#ff0000",
                command=on_delete_validated
            )
            delete_btn.pack(fill="x", pady=5)
                
        except Exception as e:
            ctk.CTkLabel(admin_window, text=f"Erreur: {e}").pack()
            logger.error(f"Erreur dashboard admin: {e}")
        finally:
            if 'conn' in locals():
                conn.close()
    
    def _open_mobile(self):
        try:
            webbrowser.open(f"http://localhost:{FLASK_PORT}")
            self._update_status("Interface mobile ouverte")
        except Exception as e:
            self._update_history(f"Erreur: {str(e)}")
            self._update_status("Erreur ouverture navigateur")
    
    def _update_history(self, message):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.history.insert(tk.END, f"[{timestamp}] {message}\n")
        self.history.see(tk.END)
    
    def _update_status(self, message):
        self.status_label.configure(text=message)
    
    def _load_history(self):
        try:
            response = requests.get(
                f'http://localhost:{FLASK_PORT}/history',
                timeout=5
            )
            response.raise_for_status()
            
            for entry in response.json():
                self._update_history(entry)
                
            self._update_status("Historique chargé")
        except requests.RequestException as e:
            self._update_history(f"Erreur chargement historique: {str(e)}")
            self._update_status("Erreur connexion historique")
        except Exception as e:
            self._update_history(f"Erreur: {str(e)}")
            self._update_status("Erreur inattendue")
    
    def _on_close(self):
        if self.scanner and self.scanner.running:
            self.scanner.stop()
        self.destroy()

# Configuration Flask
app = Flask(__name__)
CORS(app)

@app.after_request
def add_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return response

@app.route('/')
def home():
    return render_template_string(MOBILE_TEMPLATE)

@app.route('/validate', methods=['POST'])
def validate():
    try:
        data = request.get_json()
        ticket_num = data.get('ticket')
        
        if not ticket_num or not str(ticket_num).isdigit():
            return jsonify({"error": "Numéro invalide"}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Requête compatible SQLite/PostgreSQL
        if 'sqlite' in str(conn):
            cursor.execute('''
                INSERT OR REPLACE INTO tickets (ticket_number, status)
                VALUES (?, 'validé')
            ''', (ticket_num,))
        else:
            cursor.execute('''
                INSERT INTO tickets (ticket_number, status)
                VALUES (%s, 'validé')
                ON CONFLICT(ticket_number) DO UPDATE SET status='validé'
            ''', (ticket_num,))
        
        conn.commit()
        return jsonify({"message": f"Ticket {ticket_num} validé"})
    
    except Exception as e:
        logger.error(f"Erreur validation: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close()

@app.route('/verify', methods=['GET'])
def verify():
    try:
        ticket_num = request.args.get('ticket')
        if not ticket_num or not str(ticket_num).isdigit():
            return jsonify({"error": "Numéro invalide"}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT status FROM tickets WHERE ticket_number=?", (ticket_num,))
        result = cursor.fetchone()
        status = result[0] if result else 'invalide'
        return jsonify({"ticket": ticket_num, "status": status})
    
    except Exception as e:
        logger.error(f"Erreur vérification: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close()

@app.route('/export_word', methods=['GET'])
def export_word():
    try:
        doc = Document()
        doc.add_heading('Tickets Validés', 0)
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM tickets WHERE status='validé' ORDER BY timestamp DESC")
        for row in cursor.fetchall():
            doc.add_paragraph(f"Ticket {row[0]} - Validé le {row[3]}")
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='tickets_valides.docx'
        )
    
    except Exception as e:
        logger.error(f"Erreur export: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close()

@app.route('/history', methods=['GET'])
def get_history():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM tickets ORDER BY timestamp DESC LIMIT ?", (MAX_HISTORY_ENTRIES,))
        return jsonify([
            f"Ticket {row[0]} - {row[1]} - {row[3]}" 
            for row in cursor.fetchall()
        ])
    except Exception as e:
        logger.error(f"Erreur historique: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close()

@app.route('/admin', methods=['POST'])
def admin():
    try:
        data = request.get_json()
        if not data or 'password' not in data:
            return jsonify({"error": "Données manquantes"}), 400
            
        if pwd_context.verify(data['password'], ADMIN_PASSWORD_HASH):
            return jsonify({"success": True})
            
        return jsonify({"success": False, "error": "Mot de passe incorrect"}), 401
        
    except Exception as e:
        logger.error(f"Erreur admin: {e}")
        return jsonify({"error": "Erreur serveur"}), 500

@app.route('/count_validated', methods=['GET'])
def count_validated():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM tickets WHERE status='validé'")
        return jsonify({"count": cursor.fetchone()[0]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close()

@app.route('/delete_validated', methods=['POST'])
def delete_validated():
    try:
        data = request.get_json()
        if not data or 'password' not in data:
            return jsonify({"error": "Données manquantes"}), 400
            
        if not pwd_context.verify(data['password'], ADMIN_PASSWORD_HASH):
            return jsonify({"error": "Accès non autorisé"}), 401
            
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM tickets WHERE status='validé'")
        deleted_count = cursor.rowcount
        conn.commit()
    
        return jsonify({
            "success": True,
            "message": f"{deleted_count} tickets supprimés"
        })
    
    except Exception as e:
        logger.error(f"Erreur suppression: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if 'conn' in locals():
            conn.close()

@app.route('/ping')
def ping():
    return "pong", 200

def run_flask():
    """Démarre le serveur Flask"""
    logger.info(f"Démarrage sur le port {FLASK_PORT}")
    app.run(host='0.0.0.0', port=FLASK_PORT, debug=False, use_reloader=False)

if __name__ == "__main__":
    try:
        os.makedirs(QR_FOLDER, exist_ok=True)
        init_db()  # Initialise SQLite ou PostgreSQL
        
        # Démarrer Flask
        flask_thread = Thread(target=run_flask, daemon=True)
        flask_thread.start()
        
        # Démarrer l'interface desktop
        app = TicketApp()
        app._load_history()
        app.mainloop()
        
    except Exception as e:
        logger.critical(f"ERREUR: {e}")
